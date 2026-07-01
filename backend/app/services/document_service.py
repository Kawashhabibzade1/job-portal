from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

from fastapi import UploadFile

from app.models import DocumentUpdate, UploadedDocument
from app.services.profile_service import infer_profile_from_document
from app.services.serialization import model_dump
from app.services.storage import DATA_DIR, JsonStore


DOCUMENT_STORE = JsonStore("documents.json", [])
UPLOAD_DIR = DATA_DIR / "uploads"


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _document_type(filename: str, text: str) -> str:
    name = filename.lower()
    lowered = text.lower()
    if "cv" in name or "resume" in name or "curriculum vitae" in lowered:
        return "cv"
    if "certificate" in name or "diploma" in name:
        return "certificate"
    if "job" in name or "responsibilities" in lowered or "requirements" in lowered:
        return "job_description"
    return "other"


def _extract_pdf(path: Path) -> tuple[str, str]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "", "Install pypdf to extract PDF text."

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip(), ""


def _extract_docx(path: Path) -> tuple[str, str]:
    try:
        from docx import Document
    except ImportError:
        return "", "Install python-docx to extract DOCX text."

    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs).strip(), ""


def _extract_image(path: Path) -> tuple[str, str]:
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return "", "Image OCR needs pillow and pytesseract installed."

    return pytesseract.image_to_string(Image.open(path)).strip(), ""


def extract_text(path: Path, content_type: str) -> tuple[str, str, str]:
    suffix = path.suffix.lower()
    try:
        if suffix == ".txt" or content_type.startswith("text/"):
            return path.read_text(encoding="utf-8", errors="ignore").strip(), "processed", ""
        if suffix == ".pdf" or content_type == "application/pdf":
            text, message = _extract_pdf(path)
            return text, "processed" if text else "partial", message
        if suffix == ".docx":
            text, message = _extract_docx(path)
            return text, "processed" if text else "partial", message
        if content_type.startswith("image/") or suffix in {".png", ".jpg", ".jpeg", ".webp"}:
            text, message = _extract_image(path)
            return text, "processed" if text else "partial", message
    except Exception as exc:
        return "", "error", str(exc)
    return "", "unsupported", "Unsupported document type for text extraction."


def list_documents() -> list[UploadedDocument]:
    return [UploadedDocument(**item) for item in DOCUMENT_STORE.read()]


def get_document(document_id: str) -> UploadedDocument | None:
    for document in list_documents():
        if document.id == document_id:
            return document
    return None


def update_document(document_id: str, payload: DocumentUpdate) -> UploadedDocument | None:
    documents = list_documents()
    for index, document in enumerate(documents):
        if document.id != document_id:
            continue
        data = model_dump(document)
        for key, value in model_dump(payload, exclude_unset=True).items():
            if value is not None:
                data[key] = value
        updated = UploadedDocument(**data)
        documents[index] = updated
        DOCUMENT_STORE.write([model_dump(item) for item in documents])
        if updated.text:
            infer_profile_from_document(updated)
        return updated
    return None


async def store_upload(file: UploadFile) -> UploadedDocument:
    document_id = str(uuid4())
    safe_name = Path(file.filename or "upload").name
    path = UPLOAD_DIR / f"{document_id}-{safe_name}"
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    path.write_bytes(await file.read())

    content_type = file.content_type or ""
    text, status, message = extract_text(path, content_type)
    document = UploadedDocument(
        id=document_id,
        filename=safe_name,
        content_type=content_type,
        document_type=_document_type(safe_name, text),
        text=text,
        status=status,
        message=message,
        created_at=_now(),
    )
    documents = [model_dump(item) for item in list_documents()]
    documents.insert(0, model_dump(document))
    DOCUMENT_STORE.write(documents)
    if document.text:
        infer_profile_from_document(document)
    return document
